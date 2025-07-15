import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO

st.set_page_config(page_title="Treasury Report Generator", page_icon="📄")

# Begrüßung und Intro
st.title("📊 Treasury KI-Assistent – Täglicher Depot A Report")
st.markdown("""
Willkommen im Treasury-Analyse-Tool. Diese App unterstützt dich bei der tagesaktuellen Analyse deines Depot A – inkl.:
- DV01-Berechnung
- Duration
- Spread- und Liquiditätsrisiken
- Emittentenanalyse
- Automatischer Word-Export für Vorstand & ALM
""")

st.info("⬆️ Bitte lade eine Excel-Datei im Depot A Format hoch (z. B. BondsKIshort.xlsx), um die Analyse zu starten.")

uploaded_file = st.file_uploader("📤 Excel-Datei auswählen", type=["xlsx", "xls"])

# Lade Demo-Datei, falls kein Upload erfolgt ist
if uploaded_file:
    source = uploaded_file
    st.success("✅ Datei erfolgreich hochgeladen.")
else:
    st.warning("⚠️ Keine Datei hochgeladen – es werden Beispieldaten geladen.")
    source = "BondsKIshort.xlsx"

try:
    with st.spinner("🔄 Analyse läuft – bitte warten..."):
        df_raw = pd.read_excel(source, sheet_name=0, header=None)
        df_data = df_raw[1:].copy()
        df_data.columns = df_raw.iloc[0]
        df_data.reset_index(drop=True, inplace=True)

        for col in ["DV01", "Modified Duration", "Latest Yield", "3M Carry (bps)", "PP Swap Spread", "PP Govt Spread", "Size in Billions"]:
            df_data[col] = pd.to_numeric(df_data[col], errors="coerce")

        df_data["Emittent"] = df_data["Bond"].astype(str).str.extract(r'([A-Z]{2,})')

        dv01_total = df_data["DV01"].sum()
        duration_avg = df_data["Modified Duration"].mean()
        carry_avg = df_data["3M Carry (bps)"].mean()
        spread_risiko_negativ = df_data[df_data["PP Swap Spread"] < -10]
        illiquide_titel = df_data[df_data["Size in Billions"] < 1]
        exposure_emittenten = df_data.groupby("Emittent")["DV01"].sum().sort_values(ascending=False)

        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Verdana'
        font.size = Pt(10)

        doc.add_heading("Täglicher Treasury-Report – Analyse des Depot A", level=1)
        doc.add_paragraph("Datum: 15. Juli 2025")
        doc.add_paragraph("Berichtszeitraum: Tägliche Positionsbewertung")
        doc.add_paragraph("Quelle: Hochgeladene Datei" if uploaded_file else "Quelle: Beispieldatei")

        doc.add_heading("1. Gesamtüberblick – Portfolioausrichtung & Zinsrisiken", level=2)
        doc.add_paragraph(
            f"Das Portfolio weist ein aggregiertes DV01 von {dv01_total:.1f} EUR auf. "
            f"Die durchschnittliche modifizierte Duration beträgt {duration_avg:.2f} Jahre."
        )

        doc.add_heading("2. Ertragsbeiträge: Carry", level=2)
        doc.add_paragraph(
            f"Der durchschnittliche 3M Carry beträgt {carry_avg:.2f} Basispunkte über alle Positionen."
        )

        doc.add_heading("3. Spread- & Bewertungsrisiken", level=2)
        doc.add_paragraph(
            f"{len(spread_risiko_negativ)} Titel haben einen Swap Spread < –10 bps und weisen erhöhte Bewertungsrisiken auf."
        )

        doc.add_heading("4. Marktilliquidität", level=2)
        doc.add_paragraph(
            f"{len(illiquide_titel)} Positionen haben ein Emissionsvolumen unter 1 Mrd. EUR und gelten als illiquide."
        )

        doc.add_heading("5. Emittentenexposure", level=2)
        doc.add_paragraph("Top 5 Emittenten nach DV01:")

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Emittent'
        hdr_cells[1].text = 'DV01 (EUR)'
        for emittent, value in exposure_emittenten.head().items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(emittent)
            row_cells[1].text = f"{value:.2f}"

        doc.add_heading("6. Empfehlungen", level=2)
        doc.add_paragraph("- Spreadrisiken überwachen")
        doc.add_paragraph("- Illiquide Titel analysieren")
        doc.add_paragraph("- Carry validieren")

        word_io = BytesIO()
        doc.save(word_io)
        word_io.seek(0)

    st.success("✅ Analyse abgeschlossen – bereit zum Download.")
    st.download_button(
        label="📥 Word-Report herunterladen",
        data=word_io,
        file_name="Treasury_Report_DepotA.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

except Exception as e:
    st.error(f"❌ Es ist ein Fehler aufgetreten: {e}")
    st.stop()
